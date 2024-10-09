import json
import os
import re
import threading
import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox, simpledialog, ttk

import pyodbc
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tqdm.auto import tqdm
import pymysql
from pymysql.cursors import DictCursor

# 定义配置文件路径
json_file_name = "DocToTable_Config.json"


def analyze_table_sqlserver(table, table_comment):
    table_name = ""
    columns_definitions = []
    newline = ",\n    "
    primary_keys = []
    # 处理特殊符号
    table_comment = table_comment.replace("'", "''")

    # 解析SQL表和列定义
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]

        # 假定第一行包含列的注释，在第二列
        if i == 0:
            table_name = cells[1]

        # 假定列定义从第三行开始
        elif i > 1:
            column_name = cells[1]
            if not column_name:  # 如果列名为空，则跳过该行
                continue

            # 建立列定义字符串
            data_type = cells[2]
            is_null = "NULL" if cells[3] == "" else "NOT NULL"

            is_key = cells[4]
            if is_key == "主键":
                primary_keys.append(column_name)

            default_val = cells[5]
            default = ""
            auto_increment = ""
            if default_val == "自增":
                auto_increment = f"identity(1,1)"
            else:
                default = f"DEFAULT {cells[5]}" if cells[5] else ""

            column_definition = f"{column_name} {data_type} {is_null} {default} {auto_increment}".strip()
            columns_definitions.append(column_definition)

    primary_key_definition = f"PRIMARY KEY ({', '.join(primary_keys)})" if primary_keys else ""
    columns_definitions.append(primary_key_definition)

    # 组装CREATE TABLE语句
    table_sql = f"if not exists (select * from sys.sysobjects where name='{table_name}') \nbegin \n"
    table_sql += f"CREATE TABLE {table_name} (\n    {''.join({newline}).join(col for col in columns_definitions if col)}\n);"
    # table_sql = f"CREATE TABLE {table_name.lower()} ({newline}    {',{newline}    '.join(col for col in columns_definitions if col)}{newline});"

    # 添加表的注释（在MySQL中的语法示例，根据你的数据库类型可能有所不同）
    comment_sql = f"EXEC sys.sp_addextendedproperty @name=N'MS_Description', @value=N'{table_comment}', @level0type=N'Schema', @level0name=N'dbo', @level1type=N'Table', @level1name=N'{table_name} '; "
    number_of_columns = len([col for col in columns_definitions if col and not col.startswith("PRIMARY KEY")])

    column_comments = []
    for i in range(2, number_of_columns + 2):
        column_name = table.cell(i, 1).text.strip()
        column_comment = table.cell(i, 6).text.strip()
        column_comment = column_comment.replace("'", "''")
        if column_comment:
            comment_statement = f"EXEC sys.sp_addextendedproperty @name=N'MS_Description', @value=N'{column_comment}', @level0type=N'Schema', @level0name=N'dbo', @level1type=N'Table', @level1name=N'{table_name}', @level2type=N'Column', @level2name=N'{column_name}'; "
            column_comments.append(comment_statement)

    return table_sql + "\n" + comment_sql + "\n" + "\n".join(column_comments) + "\nend"


def analyze_table_mysql(table, table_comment):
    table_name = ""
    columns_definitions = []
    newline = ",\n    "
    primary_keys = []
    # 处理特殊符号
    table_comment = table_comment.replace("'", "\\'")

    # 解析SQL表和列定义
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]

        # 假定第一行包含列的注释，在第二列
        if i == 0:
            table_name = cells[1]

        # 假定列定义从第三行开始
        elif i > 1:
            column_name = cells[1]
            if not column_name:  # 如果列名为空，则跳过该行
                continue

            # 建立列定义字符串
            data_type = cells[2]
            is_null = "NULL" if cells[3] == "" else "NOT NULL"
            is_key = cells[4]

            if is_key == "主键":
                primary_keys.append(column_name)

            default_val = cells[5]
            default = ""
            auto_increment = ""
            if default_val == "自增":
                auto_increment = f"AUTO_INCREMENT"
            else:
                default = f"DEFAULT {cells[5]}" if cells[5] else ""

            comment_text = cells[6].replace("'", "\\'") if cells[6] else ""
            column_comment = f"COMMENT '{comment_text}'" if cells[6] else ""
            column_definition = f"{column_name} {data_type} {is_null} {default} {auto_increment} {column_comment}".strip()

            columns_definitions.append(column_definition)

    primary_key_definition = f"PRIMARY KEY ({', '.join(primary_keys)})" if primary_keys else ""
    if primary_key_definition:
        columns_definitions.append(primary_key_definition)

    # 组装CREATE TABLE语句
    table_sql = f"CREATE TABLE IF NOT EXISTS {table_name} (\n    {''.join({newline}).join(col for col in columns_definitions if col)}\n)"
    table_sql += f" COMMENT='{table_comment}';"

    return table_sql


def get_text_from_elem(elem):
    text = []
    for child in elem.iterchildren():
        if child.tag == qn('w:r'):  # 查找包含文本的r元素
            for subchild in child.iterchildren():
                if subchild.tag == qn('w:t'):
                    t_text = subchild.text
                    if t_text:
                        text.append(t_text)
    return ''.join(text)


def get_table_preceding_paragraph(table):
    tbl_element = table._element
    prev_element = tbl_element.getprevious()

    while prev_element is not None:
        if prev_element.tag == qn('w:p'):
            return get_text_from_elem(prev_element)
        prev_element = prev_element.getprevious()
    return ""  # 如果之前没有段落，则返回空字符串


def open_docx():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx"), ("All files", "*.*")])
    if not file_path:
        return
    else:
        open_progressbar_window2(parse_docx, file_path)


def parse_docx(progressbar, newWindow, file_path):
    dbType = db_type_combobox.get()
    try:
        doc = Document(file_path)
        sql_output = ""  # Store all the SQL statements here
        tables = doc.tables
        for i, table in enumerate(tables):
            table_comment = get_table_preceding_paragraph(table)
            # 截取空格后面的内容，1表示只分割一次
            split_comment = table_comment.split(' ', 1)
            if len(split_comment) > 1:
                table_comment = split_comment[1]

            if dbType == "sqlserver":
                sql_output += analyze_table_sqlserver(table, table_comment) + "\n\n"
            elif dbType == "mysql":
                sql_output += analyze_table_mysql(table, table_comment) + "\n\n"

            progress = int(i) / len(tables) * 100
            progressbar['value'] = progress  # 更新进度条的值
            newWindow.update_idletasks()  # 刷新界面
        newWindow.destroy()  # 任务完成后关闭窗口
        text_area.delete('1.0', tk.END)
        text_area.insert(tk.INSERT, sql_output)
    except Exception as e:
        newWindow.destroy()  # 异常时销毁窗口
        text_area.delete('1.0', tk.END)
        text_area.insert(tk.INSERT, f'Error: {e}')


# 假设这是一个全局变量来持有数据库连接信息
db_settings = {
    'dbType': 'sqlserver',
    'server': '',
    'database': '',
    'username': '',
    'password': '',
    'port': '',
    'table': ''
}


def set_cell_border(cell, **kwargs):
    """
    设置单元格边框
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 定义要设置的边缘类型
    edges = ('left', 'right', 'top', 'bottom', 'insideH', 'insideV')
    for edge in edges:
        edge_data = kwargs.get(edge, {})
        if edge_data:
            # 设置边缘属性
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', '000000'))  # 边缘设为黑色

            # 添加边缘到单元格的边界定义中
            tcBorders = tcPr.first_child_found_in('w:tcBorders')
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            tcBorders.append(element)


# 设置表格样式，所有边框都是黑色单线
border_kwargs = {
    'sz': 6,  # 边框粗细
    'val': 'single',  # 边框类型为单线
    'color': '000000',  # 边框颜色为黑色
}


def fetch_table_structure_sqlserver(server, database, username, password, port, table, progressbar, newWindow):
    """连接sqlserver数据库并获取表结构的函数。"""
    try:
        server = server + "," + port if port is not None and len(port) > 0 else server
        connection_string = f'DRIVER={{SQL Server}};SERVER={server};DATABASE={database};UID={username};PWD={password}'
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()

        query = """
                SELECT 
                    t.TABLE_NAME,
                    CONVERT(varchar(100),row_number() over(order by t.TABLE_NAME)) AS Num,
                    CONVERT(VARCHAR(MAX), ep.value) AS TABLE_COMMENT
                FROM 
                    INFORMATION_SCHEMA.TABLES t
                LEFT JOIN 
                    sys.extended_properties ep ON ep.major_id = OBJECT_ID(t.TABLE_SCHEMA + '.' + t.TABLE_NAME)
                    AND ep.minor_id = 0
                    AND ep.name = 'MS_Description'
                WHERE 
                    t.TABLE_TYPE = 'BASE TABLE'
                """
        if table is not None and len(table) > 0:
            query += " AND TABLE_NAME = ?"
            cursor.execute(query, (table,))
        else:
            cursor.execute(query)

        tables = cursor.fetchall()

        # 设置进度条
        with tqdm(total=len(tables), desc="Processing") as pbar:
            doc = Document()
            for table_name, Num, table_comment in tables:
                table_comment = table_comment if table_comment is not None and len(table_comment) > 0 else table_name
                doc.add_heading(f"{Num}. {table_comment}", level=2)  # 表名作为二级标题
                query = """
                    SELECT 
                        IC.COLUMN_NAME, 
                        IC.DATA_TYPE, 
                        CASE
                            WHEN IC.DATA_TYPE IN ('char', 'varchar', 'nchar', 'nvarchar', 'binary', 'varbinary')
                                THEN IC.CHARACTER_MAXIMUM_LENGTH
                            ELSE NULL 
                        END AS CHARACTER_MAXIMUM_LENGTH, 
                        CASE
                            WHEN IC.DATA_TYPE IN ('decimal', 'numeric', 'smallmoney', 'money', 'float', 'real')
                                THEN IC.NUMERIC_PRECISION
                            WHEN IC.DATA_TYPE IN ('int', 'smallint', 'tinyint', 'bigint')
                                THEN IC.DATETIME_PRECISION
                            WHEN IC.DATA_TYPE LIKE 'datetime2%'
                                THEN IC.DATETIME_PRECISION
                            ELSE NULL
                        END AS NUMERIC_PRECISION,
                        CASE
                            WHEN IC.DATA_TYPE IN ('decimal', 'numeric')
                                THEN IC.NUMERIC_SCALE
                            ELSE NULL 
                        END AS NUMERIC_SCALE,
                        IC.COLUMN_DEFAULT, 
                        IC.IS_NULLABLE,
                        CONVERT(VARCHAR(MAX), EP.value) AS COLUMN_COMMENT,
                        CASE WHEN PK.COLUMN_NAME IS NOT NULL THEN 'Yes' ELSE 'No' END AS IS_PRIMARY_KEY,
                        CASE WHEN SC.is_identity = 1 THEN 'Yes' ELSE 'No' END AS IS_AUTOINCREMENT
                    FROM 
                        INFORMATION_SCHEMA.COLUMNS AS IC
                    LEFT JOIN 
                        sys.columns AS SC
                        ON IC.COLUMN_NAME = SC.name AND OBJECT_NAME(SC.object_id) = IC.TABLE_NAME
                    LEFT JOIN 
                        sys.extended_properties AS EP
                        ON EP.major_id = SC.object_id AND EP.minor_id = SC.column_id AND EP.name = 'MS_Description'
                    LEFT JOIN 
                        (SELECT 
                            KCU.TABLE_NAME, 
                            KCU.COLUMN_NAME 
                        FROM 
                            INFORMATION_SCHEMA.KEY_COLUMN_USAGE AS KCU
                        JOIN 
                            INFORMATION_SCHEMA.TABLE_CONSTRAINTS AS TC
                        ON KCU.CONSTRAINT_NAME = TC.CONSTRAINT_NAME
                        WHERE TC.CONSTRAINT_TYPE = 'PRIMARY KEY') AS PK
                        ON IC.TABLE_NAME = PK.TABLE_NAME AND IC.COLUMN_NAME = PK.COLUMN_NAME
                    WHERE 
                        IC.TABLE_NAME = ?
                        """
                # 参数化查询以确保安全
                cursor.execute(query, (table_name,))
                columns = cursor.fetchall()

                t = doc.add_table(rows=len(columns) + 2, cols=7)
                # 遍历表格并应用边框
                for row in t.rows:
                    for cell in row.cells:
                        # 为每个单元格分别设置所有四个边框和内部边框
                        set_cell_border(cell,
                                        left=border_kwargs,
                                        right=border_kwargs,
                                        top=border_kwargs,
                                        bottom=border_kwargs,
                                        insideH=border_kwargs,
                                        insideV=border_kwargs)

                t.cell(0, 0).text = '表名'
                t.cell(0, 1).merge(t.cell(0, 2)).text = table_name
                t.cell(0, 3).merge(t.cell(0, 4)).merge(t.cell(0, 5)).text = '所属数据库'
                t.cell(0, 6).text = database

                t.cell(1, 0).text = '序号'
                t.cell(1, 1).text = '字段名'
                t.cell(1, 2).text = '字段类型'
                t.cell(1, 3).text = '非空'
                t.cell(1, 4).text = '键'
                t.cell(1, 5).text = '默认值'
                t.cell(1, 6).text = '字段说明'

                for i, column in enumerate(columns):
                    row = t.rows[i + 2]
                    row.cells[0].text = str(i + 1)
                    row.cells[1].text = column.COLUMN_NAME
                    column_type = column.DATA_TYPE
                    column_length = column.CHARACTER_MAXIMUM_LENGTH

                    if (column_type == "nvarchar" or column_type == "varchar") and str(column_length) == "-1":
                        column_length = "max"
                    if column_length is None or len(str(column_length)) < 1:
                        column_length = column.NUMERIC_PRECISION
                    if column_type == "decimal" or column_type == "numeric":
                        column_length = f"{column.NUMERIC_PRECISION},{column.NUMERIC_SCALE}"
                    if column_length:
                        column_type += f" ({column_length})"
                    row.cells[2].text = column_type
                    IS_NULLABLE = str(column.IS_NULLABLE)
                    row.cells[3].text = '是' if IS_NULLABLE == 'NO' else ''
                    IS_PRIMARY_KEY = str(column.IS_PRIMARY_KEY)
                    row.cells[4].text = '主键' if IS_PRIMARY_KEY == 'Yes' else ''
                    COLUMN_DEFAULT = str(column.COLUMN_DEFAULT) if column.COLUMN_DEFAULT else ''
                    COLUMN_DEFAULT = COLUMN_DEFAULT.strip("(())")
                    row.cells[5].text = COLUMN_DEFAULT
                    IS_AUTOINCREMENT = str(column.IS_AUTOINCREMENT)
                    if IS_AUTOINCREMENT == 'Yes':
                        row.cells[5].text = '自增'
                    row.cells[6].text = column.COLUMN_COMMENT if column.COLUMN_COMMENT else ''
                # 更新进度条
                pbar.update(1)
                progress = int(Num) / len(tables) * 100
                progressbar['value'] = progress  # 更新进度条的值
                newWindow.update_idletasks()  # 刷新界面
            newWindow.destroy()  # 任务完成后关闭窗口
        return doc

    except Exception as e:
        newWindow.destroy()  # 出现异常时关闭窗口
        messagebox.showerror("Error", f"Could not connect to database: {e}")
        return None


def fetch_table_structure_mysql(server, database, username, password, port, table, progressbar, newWindow):
    """连接mysql数据库并获取表结构的函数。"""
    try:
        conn = pymysql.connect(host=server,
                               user=username,
                               password=password,
                               database=database,
                               port=int(port) if port else 3306,
                               cursorclass=DictCursor)

        with (conn):
            with conn.cursor() as cursor:
                # 获取表的查询，包括表注释
                query = """
                        SELECT 
                            t.TABLE_NAME,
                            (@rownum:=@rownum+1) AS Num,
                            t.TABLE_COMMENT
                        FROM 
                            INFORMATION_SCHEMA.TABLES t
                        JOIN (SELECT @rownum:=0) r  # 初始化行号
                        WHERE 
                            t.TABLE_SCHEMA = %s
                            AND t.TABLE_TYPE = 'BASE TABLE'
                        """
                parameters = [database]
                if table:
                    query += " AND t.TABLE_NAME = %s"
                    parameters.append(table)

                cursor.execute(query, parameters)
                tables = cursor.fetchall()

                # 设置进度条
                with tqdm(total=len(tables), desc="Processing") as pbar:
                    doc = Document()
                for table_info in tables:
                    table_name = table_info["TABLE_NAME"]
                    Num = table_info["Num"]
                    table_comment = table_info["TABLE_COMMENT"]

                    table_comment = table_comment if table_comment is not None and len(
                        table_comment) > 0 else table_name
                    doc.add_heading(f"{Num}. {table_comment}", level=2)  # 表名作为二级标题

                    # 获取列的查询，包括列注释和是否为主键
                    query_columns = """
                        SELECT 
                            c.COLUMN_NAME, 
                            c.DATA_TYPE, 
                            c.CHARACTER_MAXIMUM_LENGTH, 
                            c.NUMERIC_PRECISION, 
                            c.NUMERIC_SCALE, 
                            c.COLUMN_DEFAULT, 
                            c.IS_NULLABLE,
                            IF(k.COLUMN_NAME IS NOT NULL, 'Yes', 'No') AS IS_PRIMARY_KEY,
                            c.COLUMN_COMMENT,
                            IF(t.AUTO_INCREMENT IS NOT NULL AND c.EXTRA = 'auto_increment', 'Yes', 'No') AS IS_AUTO_INCREMENT
                        FROM 
                            INFORMATION_SCHEMA.COLUMNS c
                        LEFT JOIN 
                            INFORMATION_SCHEMA.KEY_COLUMN_USAGE k ON c.TABLE_SCHEMA = k.TABLE_SCHEMA
                            AND c.TABLE_NAME = k.TABLE_NAME
                            AND c.COLUMN_NAME = k.COLUMN_NAME
                            AND k.CONSTRAINT_NAME = 'PRIMARY'
                        LEFT JOIN 
                            INFORMATION_SCHEMA.TABLES t ON c.TABLE_SCHEMA = t.TABLE_SCHEMA
                            AND c.TABLE_NAME = t.TABLE_NAME
                            AND t.TABLE_TYPE = 'BASE TABLE'
                        WHERE 
                            c.TABLE_SCHEMA = %s
                            AND c.TABLE_NAME = %s
                        ORDER BY 
                            c.ORDINAL_POSITION
                    """
                    cursor.execute(query_columns, (database, table_name))
                    columns = cursor.fetchall()

                    t = doc.add_table(rows=len(columns) + 2, cols=7)
                    # 遍历表格并应用边框
                    for row in t.rows:
                        for cell in row.cells:
                            # 为每个单元格分别设置所有四个边框和内部边框
                            set_cell_border(cell,
                                            left=border_kwargs,
                                            right=border_kwargs,
                                            top=border_kwargs,
                                            bottom=border_kwargs,
                                            insideH=border_kwargs,
                                            insideV=border_kwargs)

                    t.cell(0, 0).text = '表名'
                    t.cell(0, 1).merge(t.cell(0, 2)).text = table_name
                    t.cell(0, 3).merge(t.cell(0, 4)).merge(t.cell(0, 5)).text = '所属数据库'
                    t.cell(0, 6).text = database

                    t.cell(1, 0).text = '序号'
                    t.cell(1, 1).text = '字段名'
                    t.cell(1, 2).text = '字段类型'
                    t.cell(1, 3).text = '非空'
                    t.cell(1, 4).text = '键'
                    t.cell(1, 5).text = '默认值'
                    t.cell(1, 6).text = '字段说明'

                    for i, column in enumerate(columns):
                        row = t.rows[i + 2]
                        row.cells[0].text = str(i + 1)
                        row.cells[1].text = column["COLUMN_NAME"]
                        column_type = column["DATA_TYPE"]
                        column_length = column["CHARACTER_MAXIMUM_LENGTH"]

                        if (column_type == "nvarchar" or column_type == "varchar") and str(column_length) == "-1":
                            column_length = "max"
                        if column_length is None or len(str(column_length)) < 1:
                            column_length = column["NUMERIC_PRECISION"]
                        if column_type == "decimal" or column_type == "numeric":
                            column_length = f"{column['NUMERIC_PRECISION']},{column['NUMERIC_SCALE']}"
                        if column_length:
                            column_type += f" ({column_length})"
                        row.cells[2].text = column_type
                        IS_NULLABLE = str(column["IS_NULLABLE"])
                        row.cells[3].text = '是' if IS_NULLABLE == 'NO' else ''
                        IS_PRIMARY_KEY = str(column["IS_PRIMARY_KEY"])
                        row.cells[4].text = '主键' if IS_PRIMARY_KEY == 'Yes' else ''
                        row.cells[5].text = str(column["COLUMN_DEFAULT"]) if column["COLUMN_DEFAULT"] else ''
                        IS_AUTO_INCREMENT = str(column["IS_AUTO_INCREMENT"])
                        if IS_AUTO_INCREMENT == 'Yes':
                            row.cells[5].text = '自增'
                        row.cells[6].text = column["COLUMN_COMMENT"] if column["COLUMN_COMMENT"] else ''
                    # 更新进度条
                    pbar.update(1)
                    progress = int(Num) / len(tables) * 100
                    progressbar['value'] = progress  # 更新进度条的值
                    newWindow.update_idletasks()  # 刷新界面
            newWindow.destroy()  # 任务完成后关闭窗口
        return doc

    except Exception as e:
        newWindow.destroy()  # 出现异常时关闭窗口
        messagebox.showerror("Error", f"Could not connect to database: {e}")
        return None


def update_config_json(new_config):
    # 写入内容到新文件中
    with open(json_file_name, 'w', encoding='utf-8') as file:
        json.dump(new_config, file, ensure_ascii=False, indent=4)


def convert_to_word(progressbar, newWindow):
    # 从输入字段获取数据库连接信息
    dbType = db_type_combobox.get()
    server = server_entry.get()
    database = database_entry.get()
    username = username_entry.get()
    password = password_entry.get()
    port = port_entry.get()
    table = table_entry.get()
    if server is None or len(server) < 1:
        messagebox.showerror("Error", "数据库链接地址server不能为空")
        newWindow.destroy()
        return
    if database is None or len(database) < 1:
        messagebox.showerror("Error", "数据库名称database不能为空")
        newWindow.destroy()
        return
    if username is None or len(username) < 1:
        messagebox.showerror("Error", "数据库登录名username不能为空")
        newWindow.destroy()
        return
    if password is None or len(password) < 1:
        messagebox.showerror("Error", "数据库链接密码password不能为空")
        newWindow.destroy()
        return

    new_config = {
        'dbType': dbType,
        'server': server,
        'database': database,
        'username': username,
        'password': password,
        'port': port,
        'table': table
    }
    update_config_json(new_config)

    if dbType == "sqlserver":
        doc = fetch_table_structure_sqlserver(server, database, username, password, port, table, progressbar, newWindow)
    elif dbType == "mysql":
        doc = fetch_table_structure_mysql(server, database, username, password, port, table, progressbar, newWindow)

    if doc:
        filepath = filedialog.asksaveasfilename(defaultextension=".docx")
        if filepath:
            doc.save(filepath)
            messagebox.showinfo("Info", f"The Word document was created successfully: {filepath}")


def load_db_settings():
    # 检查配置文件是否存在
    if not os.path.isfile(json_file_name):
        return db_settings

    # 加载配置文件
    with open(json_file_name, 'r') as config_file:
        return json.load(config_file)


# 配置文件的数据库设置
db_settings = load_db_settings()

# 创建窗体
root = tk.Tk()
root.title("Word文档与表结构转换工具")

# 数据库类型选项
db_options = ["sqlserver", "mysql"]
# 创建一个Label
label_db_type = tk.Label(root, text="DbType:")
label_db_type.pack()
# 创建一个Combobox，包含数据库类型选项
db_type_combobox = ttk.Combobox(root, values=db_options, width=17)
db_type_combobox.pack()
db_type_combobox.set(db_settings['dbType'])  # 设置初始显示文本

# 创建标签和输入字段
label_server = tk.Label(root, text="Server:")
label_server.pack()
server_entry = tk.Entry(root)
server_entry.pack()
server_entry.insert(0, db_settings['server'])

label_database = tk.Label(root, text="Database:")
label_database.pack()
database_entry = tk.Entry(root)
database_entry.pack()
database_entry.insert(0, db_settings['database'])

label_username = tk.Label(root, text="Username:")
label_username.pack()
username_entry = tk.Entry(root)
username_entry.pack()
username_entry.insert(0, db_settings['username'])

label_password = tk.Label(root, text="Password:")
label_password.pack()
password_entry = tk.Entry(root, show="*")
password_entry.pack()
password_entry.insert(0, db_settings['password'])

label_port = tk.Label(root, text="Port:")
label_port.pack()
port_entry = tk.Entry(root)
port_entry.pack()
port_entry.insert(0, db_settings['port'])

label_table = tk.Label(root, text="table:")
label_table.pack()
table_entry = tk.Entry(root)
table_entry.pack()
table_entry.insert(0, db_settings['table'])


def open_progressbar_window(fun):
    newWindow = tk.Toplevel(root)
    newWindow.title("执行进度")
    newWindow.geometry("300x50")
    progressbar = ttk.Progressbar(newWindow, orient=tk.HORIZONTAL, length=280, mode='determinate')
    progressbar.pack(pady=10)

    # 在单独的线程中开始任务
    threading.Thread(target=fun, args=(progressbar, newWindow,)).start()


def open_progressbar_window2(fun, param):
    newWindow = tk.Toplevel(root)
    newWindow.title("执行进度")
    newWindow.geometry("300x50")
    progressbar = ttk.Progressbar(newWindow, orient=tk.HORIZONTAL, length=280, mode='determinate')
    progressbar.pack(pady=10)

    # 在单独的线程中开始任务
    threading.Thread(target=fun, args=(progressbar, newWindow, param,)).start()


open_db_button = tk.Button(root, text="数据库表结构保存为Word文档",
                           command=lambda: open_progressbar_window(convert_to_word))
open_db_button.pack(side=tk.TOP, pady=10)

text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD)
text_area.pack(side=tk.TOP, pady=10)

open_button = tk.Button(root, text="Word文档转表结构脚本", command=open_docx)
open_button.pack(side=tk.TOP, pady=10)

root.mainloop()
