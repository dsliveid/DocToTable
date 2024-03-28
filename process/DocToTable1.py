from docx import Document
import os

# 替换为你的Word文档路径，建议使用原始字符串
doc_path = r'D:\表结构测试.docx'  # 确保是.docx格式

# 检查文件是否存在
if not os.path.exists(doc_path):
    print(f'文件"{doc_path}"不存在。')
elif not doc_path.endswith('.docx'):
    print('文件不是.docx格式。请开启Word并将文件另存为.docx格式后重试。')
else:
    try:
        # 加载Word文档
        doc = Document(doc_path)

        # 输出段落
        print('文档段落：')
        for i, para in enumerate(doc.paragraphs):
            print(f'段落{i + 1}: {para.text}')

        # 输出表格内容
        print('\n文档表格：')
        for i, table in enumerate(doc.tables):
            print(f'表格{i + 1}:')
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text, end=' | ')
                print()  # 新的一行

        # 输出图片说明，但不能直接从docx文档中提取图片内容
        print('\n文档中的图片：')
        for i, shape in enumerate(doc.inline_shapes):
            # inline_shapes可以包含图片，但python-docx无法直接访问图片二进制数据
            print(f'图片{i + 1}: 图片可能在文档中，但python-docx暂时不能提取它们。')

        # 注意，python-docx库目前还不支持提取原始图片

    except Exception as e:
        print(f'处理Word文档时发生错误：{e}')
