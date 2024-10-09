import json
import os
import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox, simpledialog

import pyodbc
from docx import Document


def analyze_table(table, table_comment):
    table_name = ""
    db_name = ""
    columns_definitions = []
    newline = ",\n    "
    primary_keys = []

    # 解析SQL表和列定义
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]

        # 假定第一行包含列的注释，在第二列
        if i == 0:
            table_name = cells[1]
            db_name = cells[6]

        # 假定列定义从第三行开始
        elif i > 1:
            column_name = cells[1]
            if not column_name:  # 如果列名为空，则跳过该行
                continue

            # 建立列定义字符串
            data_type = cells[2]
            is_null = "NULL" if cells[3] == "" else "NOT NULL"
            is_key = cells[4]
            default = f"DEFAULT {cells[5]}" if cells[5] else ""

            column_definition = f"{column_name} {data_type} {is_null} {default}".strip()

            if is_key == "主键":
                primary_keys.append(column_name)
                # 在MySQL中，设置了PRIMARY KEY的列不能被定义为NULL
                # column_definition = column_definition.replace("NULL", "")
            elif is_key == "外键":
                # 在此示例中，外键的处理略过了，因为要建立外键，还需要知道外键引用了哪个表和列
                pass

            columns_definitions.append(column_definition)

    primary_key_definition = f"PRIMARY KEY ({', '.join(primary_keys)})" if primary_keys else ""
    columns_definitions.append(primary_key_definition)

    # 组装CREATE TABLE语句
    table_sql = f"if not exits (select * from sys.sysobjects where name='{table_name}') \nbegin \n"
    table_sql += f"CREATE TABLE {table_name} (\n    {''.join({newline}).join(col for col in columns_definitions if col)}\n);"
    # table_sql = f"CREATE TABLE {table_name.lower()} ({newline}    {',{newline}    '.join(col for col in columns_definitions if col)}{newline});"

    # 添加表的注释（在MySQL中的语法示例，根据你的数据库类型可能有所不同）
    comment_sql = f"EXEC sys.sp_addextendedproperty @name=N'MS_Description', @value=N'{table_comment}', @level0type=N'Schema', @level0name=N'dbo', @level1type=N'Table', @level1name=N'{table_name} '; "
    number_of_columns = len([col for col in columns_definitions if col and not col.startswith("PRIMARY KEY")])

    column_comments = []
    for i in range(2, number_of_columns + 2):
        column_name = table.cell(i, 1).text.strip()
        column_comment = table.cell(i, 6).text.strip()
        if column_comment:
            comment_statement = f"EXEC sys.sp_addextendedproperty @name=N'MS_Description', @value=N'{column_comment}', @level0type=N'Schema', @level0name=N'dbo', @level1type=N'Table', @level1name=N'{table_name}', @level2type=N'Column', @level2name=N'{column_name}'; "
            column_comments.append(comment_statement)

    return table_sql + "\n" + comment_sql + "\n" + "\n".join(column_comments) + "\nend"


def get_table_preceding_paragraph(table):
    tbl_element = table._element
    prev_element = tbl_element.getprevious()

    if prev_element is not None and prev_element.tag.endswith('p'):
        return prev_element.text
    else:
        # 找到真正的文本段落
        while prev_element is not None and not prev_element.tag.endswith('p'):
            prev_element = prev_element.getprevious()
        if prev_element is not None:
            return prev_element.text
    return ""  # 如果之前没有段落，则返回空字符串


def open_docx():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx"), ("All files", "*.*")])
    if not file_path:
        return

    try:
        doc = Document(file_path)
        sql_output = ""  # Store all the SQL statements here
        for table in doc.tables:
            table_comment = get_table_preceding_paragraph(table)
            sql_output += analyze_table(table, table_comment) + "\n\n"

        text_area.delete('1.0', tk.END)
        text_area.insert(tk.INSERT, sql_output)
    except Exception as e:
        text_area.delete('1.0', tk.END)
        text_area.insert(tk.INSERT, f'Error: {e}')


# 假设这是一个全局变量来持有数据库连接信息
db_settings = {
    'server': '',
    'database': '',
    'username': '',
    'password': '',
    'port': '',
    'table': ''
}


def fetch_table_structure(server, database, username, password, port, table):
    """连接数据库并获取表结构的函数。"""
    try:
        connection_string = f'DRIVER={{SQL Server}};SERVER={server},{port};DATABASE={database};UID={username};PWD={password}'
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()

        query = """
                SELECT 
                    t.TABLE_NAME,
                    t.TABLE_TYPE,
                    CONVERT(VARCHAR(MAX), ep.value) AS TABLE_COMMENT
                FROM 
                    INFORMATION_SCHEMA.TABLES t
                LEFT JOIN 
                    sys.extended_properties ep ON ep.major_id = OBJECT_ID(t.TABLE_SCHEMA + '.' + t.TABLE_NAME)
                    AND ep.minor_id = 0
                    AND ep.name = 'MS_Description'
                WHERE 
                    t.TABLE_TYPE = 'BASE TABLE'
                """
        if table is not None and len(table) > 0:
            query += " AND TABLE_NAME = ?"
            cursor.execute(query, (table,))
        else:
            cursor.execute(query)

        tables = cursor.fetchall()

        doc = Document()
        for table_name, table_type, table_comment in tables:
            doc.add_heading(f"{table_comment}", level=2)  # 表名作为二级标题
            query = """
                    SELECT COLUMN_NAME, DATA_TYPE, CHARACTER_MAXIMUM_LENGTH, COLUMN_DEFAULT, IC.IS_NULLABLE,
                        CONVERT(VARCHAR(MAX), EP.value) AS COLUMN_COMMENT
                    FROM INFORMATION_SCHEMA.COLUMNS AS IC
                    LEFT JOIN sys.columns AS SC
                        ON IC.COLUMN_NAME = SC.name AND OBJECT_NAME(SC.object_id) = IC.TABLE_NAME
                    LEFT JOIN sys.extended_properties AS EP
                        ON EP.major_id = SC.object_id AND EP.minor_id = SC.column_id AND EP.name = 'MS_Description'
                    WHERE IC.TABLE_NAME = ?
                    """  # 参数化查询以确保安全
            cursor.execute(query, (table_name,))
            columns = cursor.fetchall()

            t = doc.add_table(rows=len(columns) + 2, cols=5)
            t.cell(0, 0).text = '表名'
            t.cell(0, 1).text = table_name
            t.cell(0, 2).merge(t.cell(0, 3)).text = '所属数据库'
            t.cell(0, 4).text = database

            t.cell(1, 0).text = '序号'
            t.cell(1, 1).text = '字段名'
            t.cell(1, 2).text = '字段类型'
            t.cell(1, 3).text = '默认值'
            t.cell(1, 4).text = '备注'

            for i, column in enumerate(columns):
                row = t.rows[i + 2]
                row.cells[0].text = str(i + 1)
                row.cells[1].text = column.COLUMN_NAME
                column_type = column.DATA_TYPE
                if column.CHARACTER_MAXIMUM_LENGTH:
                    column_type += f" ({column.CHARACTER_MAXIMUM_LENGTH})"
                row.cells[2].text = column_type
                row.cells[3].text = str(column.COLUMN_DEFAULT) if column.COLUMN_DEFAULT else ''
                row.cells[4].text = column.COLUMN_COMMENT if column.COLUMN_COMMENT else ''

        return doc

    except Exception as e:
        messagebox.showerror("Error", f"Could not connect to database: {e}")
        return None

def convert_to_word():
    # 从输入字段获取数据库连接信息
    server = server_entry.get()
    database = database_entry.get()
    username = username_entry.get()
    password = password_entry.get()
    port = port_entry.get()
    table = table_entry.get()

    doc = fetch_table_structure(server, database, username, password, port, table)
    if doc:
        filepath = filedialog.asksaveasfilename(defaultextension=".docx")
        if filepath:
            doc.save(filepath)
            messagebox.showinfo("Info", f"The Word document was created successfully: {filepath}")


def load_db_settings():
    # 检查配置文件是否存在
    if not os.path.isfile('../DocToTable_Config_mysql.json'):
        return db_settings

    # 加载配置文件
    with open('../DocToTable_Config_mysql.json', 'r') as config_file:
        return json.load(config_file)

# 配置文件的数据库设置
db_settings = load_db_settings()

# 创建窗体
root = tk.Tk()
root.title("Word文档与表结构转换工具")

# 创建标签和输入字段
label_server = tk.Label(root, text="Server:")
label_server.pack()
server_entry = tk.Entry(root)
server_entry.pack()
server_entry.insert(0, db_settings['server'])

label_database = tk.Label(root, text="Database:")
label_database.pack()
database_entry = tk.Entry(root)
database_entry.pack()
database_entry.insert(0, db_settings['database'])

label_username = tk.Label(root, text="Username:")
label_username.pack()
username_entry = tk.Entry(root)
username_entry.pack()
username_entry.insert(0, db_settings['username'])

label_password = tk.Label(root, text="Password:")
label_password.pack()
password_entry = tk.Entry(root, show="*")
password_entry.pack()
password_entry.insert(0, db_settings['password'])

label_port = tk.Label(root, text="Port:")
label_port.pack()
port_entry = tk.Entry(root)
port_entry.pack()
port_entry.insert(0, db_settings['port'])

label_table = tk.Label(root, text="table:")
label_table.pack()
table_entry = tk.Entry(root)
table_entry.pack()
table_entry.insert(0, db_settings['table'])

text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD)
text_area.pack(side=tk.TOP, pady=10)

open_button = tk.Button(root, text="打开Word文档", command=open_docx)
open_button.pack(side=tk.TOP, pady=10)

open_db_button = tk.Button(root, text="表结构转Word文档", command=convert_to_word)
open_db_button.pack(side=tk.TOP, pady=10)

root.mainloop()
