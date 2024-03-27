import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document


def open_docx():
    # 使用文件对话框选择.docx文件
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx"), ("All files", "*.*")])
    if not file_path:
        return

    try:
        # 加载Word文档
        doc = Document(file_path)

        # 清除文本区域
        text_area.delete('1.0', tk.END)

        # 将解析结果输出到文本区域
        text_area.insert(tk.END, '文档段落：\n')
        for para in doc.paragraphs:
            text_area.insert(tk.END, f'{para.text}\n\n')
        text_area.insert(tk.END, '文档表格：\n')

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_area.insert(tk.END, f'{cell.text}\t| ')
                text_area.insert(tk.END, '\n')
            text_area.insert(tk.END, '\n')
    except Exception as e:
        messagebox.showerror("错误", f"打开文件时出错: {e}")


# 创建主窗口
root = tk.Tk()
root.title("Word文档解析器")

# 创建菜单栏
menu_bar = tk.Menu(root)
file_menu = tk.Menu(menu_bar, tearoff=0)
file_menu.add_command(label="打开", command=open_docx)
file_menu.add_command(label="退出", command=root.quit)
menu_bar.add_cascade(label="文件", menu=file_menu)
root.config(menu=menu_bar)

# 创建滚动文本区域以显示文档内容
text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD)
text_area.pack(fill=tk.BOTH, expand=True)

# 运行主循环
root.mainloop()
